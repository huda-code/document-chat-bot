import os
import gradio as gr
from docx import Document
from sentence_transformers import SentenceTransformer, util
import numpy as np
import nltk
from nltk.tokenize import sent_tokenize
import shutil

# Download NLTK data
nltk.download('punkt')

# Initialize the sentence transformer model
model = SentenceTransformer('all-MiniLM-L6-v2')

class SimpleDirectoryReader:
    def __init__(self, directory):
        self.directory = directory
    
    def load_data(self):
        documents = []
        for filename in os.listdir(self.directory):
            if filename.endswith(".docx"):
                doc_path = os.path.join(self.directory, filename)
                documents.extend(self._read_docx(doc_path))
        return documents
    
    def _read_docx(self, filepath):
        doc = Document(filepath)
        sentences = []
        for para in doc.paragraphs:
            if para.text.strip():
                sentences.extend(sent_tokenize(para.text))
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    sentences.append(" | ".join(row_data))
        return sentences

class VectorStoreIndex:
    @classmethod
    def from_documents(cls, documents):
        index = cls()
        index.documents = documents
        index.doc_embeddings = model.encode(documents, convert_to_tensor=True)
        return index
    
    def __init__(self):
        self.documents = []
        self.doc_embeddings = None
        self.storage_context = self.StorageContext()
    
    def as_query_engine(self):
        return self.QueryEngine(self.documents, self.doc_embeddings)
    
    class StorageContext:
        def persist(self, persist_dir):
            print(f"Mock: Persisting storage to {persist_dir}")

    class QueryEngine:
        def __init__(self, documents, doc_embeddings):
            self.documents = documents
            self.doc_embeddings = doc_embeddings
        
        def query(self, query):
            response = self._search_documents(query)
            return response
        
        def _search_documents(self, query):
            query_embedding = model.encode(query, convert_to_tensor=True)
            cosine_scores = util.pytorch_cos_sim(query_embedding, self.doc_embeddings)[0]
            best_match_idx = np.argmax(cosine_scores.numpy())
            best_match_score = cosine_scores[best_match_idx]
            if best_match_score > 0.5:  # Set a threshold for relevance
                best_sentence = self.documents[best_match_idx].strip()
                surrounding_context = self._get_surrounding_sentences(best_match_idx)
                return f"Found in document: {best_sentence}\n\nContext: {surrounding_context}"
            return "Sorry, I couldn't find an answer to your question in the document."
        
        def _get_surrounding_sentences(self, idx, window=1):
            start_idx = max(0, idx - window)
            end_idx = min(len(self.documents), idx + window + 1)
            surrounding_sentences = " ".join(self.documents[start_idx:end_idx])
            return surrounding_sentences

# Set the directory for persistent storage
PERSIST_DIR = "./storage"
CHAT_LOG = "chat_log.txt"

# Clear previous chat log
if os.path.exists(CHAT_LOG):
    os.remove(CHAT_LOG)

# Function to upload and process the document
def upload_document(file):
    try:
        global index
        index = None  # Clear the previous index

        # Clear previous files in storage and docs
        if os.path.exists(PERSIST_DIR):
            shutil.rmtree(PERSIST_DIR)
        os.makedirs(PERSIST_DIR)
        if os.path.exists("docs"):
            shutil.rmtree("docs")
        os.makedirs("docs")

        # Save the uploaded file
        file_path = os.path.join("docs", "uploaded_document.docx")
        with open(file_path, "wb") as f:
            f.write(file)

        # Process the document
        documents = SimpleDirectoryReader("docs").load_data()
        index = VectorStoreIndex.from_documents(documents)
        index.storage_context.persist(persist_dir=PERSIST_DIR)

        return "Document uploaded and processed successfully."
    except Exception as e:
        return f"Error: {str(e)}"

# Initialize index
index = None

# Function to handle chatbot queries
def chatbot(query):
    if index:
        response = index.as_query_engine().query(query)
    else:
        response = "No document uploaded or processed."
    
    # Log the question and response
    with open(CHAT_LOG, "a") as f:
        f.write(f"Q: {query}\nA: {response}\n\n")

    return response

# Function to download chat log
def download_chat_log():
    return CHAT_LOG

# Create and launch the Gradio interface
iface = gr.Blocks()

with iface:
    gr.Markdown("# GPT-Based Chatbot")
    with gr.Row():
        with gr.Column():
            upload_btn = gr.File(label="Upload Document", file_types=[".txt", ".docx"], type="binary")
            upload_output = gr.Textbox(label="Upload Status")
        with gr.Column():
            chatbot_input = gr.Textbox(label="Ask a question", placeholder="Type your question here...")
            chatbot_output = gr.Textbox(label="Answer")
            submit_btn = gr.Button("Submit")
            download_btn = gr.Button("Download Chat Log")
    
    upload_btn.upload(upload_document, inputs=upload_btn, outputs=upload_output)
    submit_btn.click(chatbot, inputs=chatbot_input, outputs=chatbot_output)
    download_btn.click(download_chat_log, inputs=[], outputs=gr.File(label="Chat Log"))

iface.launch()

